'''
Editing resumes using OpenAI API

'''

# imports

import os
import requests
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from IPython.display import Markdown, display
from openai import OpenAI
from docx import Document


# summarize the content of a resume
def summarize(resume, systemPrompt=None):
    sys_prompt = "You are excellent at reading resumes and \
    making a concise summary of the resume \
    " if not systemPrompt else systemPrompt
    
    user_prompt = f"Review and summarize the content of this resume \n\n{resume}"

    messages = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt}
    ]
    response = openai.chat.completions.create(model="gpt-4o-mini", messages=messages)
    return response.choices[0].message.content

# generate an excellent resume to match job description
def generate_excellent_res(resume_summary, job_description):
    sys_prompt = "You are an excellent assistant for reading resume summary and a \
    job description, editing, and generating an excellent matching resume to the job description."
    user_prompt = f"Review the resume summary and generate an excellent edited resume more tailored \
    the job description. \
    \nResume Summary: {resume_summary} \n\
    Job description: {job_description}"
    
    messages = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt}
    ]
    response = openai.chat.completions.create(model="gpt-4o-mini", messages=messages)
    return response.choices[0].message.content


if __name__=='__main__':

    # Load environment variables in a file called .env

    load_dotenv(override=True)
    api_key = os.getenv('OPENAI_API_KEY')

    # Check the key

    if not api_key:
        print("No API key was found - please head over to the troubleshooting notebook in this folder to identify & fix!")
    elif not api_key.startswith("sk-proj-"):
        print("An API key was found, but it doesn't start sk-proj-; please check you're using the right key - see troubleshooting notebook")
    elif api_key.strip() != api_key:
        print("An API key was found, but it looks like it might have space or tab characters at the start or end - please remove them - see troubleshooting notebook")
    else:
        print("API key found and looks good so far!")


    openai = OpenAI()

    # If this doesn't work, try Kernel menu >> Restart Kernel and Clear Outputs Of All Cells, then run the cells from the top of this notebook down.
    # If it STILL doesn't work (horrors!) then please see the Troubleshooting notebook in this folder for full instructions



    current_path = os.getcwd()

    # reading resume content
    resume_file_name = input("Enter resume file name: ").strip()
    document =''
    try:
        document = Document(current_path + "/" + resume_file_name)
    except:
        raise Exception(f"Resume file '{resume_file_name}' not found or unreadable.")


    resume = ""
    for paragraph in document.paragraphs:
        resume += paragraph.text + "\n"

    # reading job description
    job_description_name = input("Enter job description file name: ").strip()
    job = ''
    try:
        job = Document(current_path + '/' + job_description_name)
    except:
        raise Exception(f"Resume file '{resume_file_name}' not found or unreadable.")

        
    job_description = ""
    for para in job.paragraphs:
        job_description += para.text + "\n"

    resume_summary = summarize(resume)
    print("\n\n")
    display(Markdown(resume_summary))
    # print("Resume Summary: \n", resume_summary)
    print("\n\n")

    cover_let = generate_excellent_res(resume_summary, job_description)
    display(Markdown(cover_let))
    print("\n\n")
